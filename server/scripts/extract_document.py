#!/usr/bin/env python3
import json
import os
import re
import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader


def clean_text(value: str) -> str:
    text = (value or "").replace("\r", "")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def read_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    parts = []
    for index, page in enumerate(reader.pages, start=1):
      text = page.extract_text() or ""
      text = clean_text(text)
      if text:
          parts.append(f"## 第 {index} 页\n\n{text}")
    return "\n\n".join(parts)


def read_docx(path: Path) -> str:
    doc = Document(str(path))
    parts = []
    for paragraph in doc.paragraphs:
        text = clean_text(paragraph.text)
        if text:
            parts.append(text)
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [clean_text(cell.text) for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            parts.append("\n".join(rows))
    return "\n\n".join(parts)


def read_xlsx(path: Path) -> str:
    workbook = load_workbook(filename=str(path), data_only=True, read_only=True)
    parts = []
    for sheet in workbook.worksheets:
        rows = [f"# 工作表：{sheet.title}"]
        for row in sheet.iter_rows(values_only=True):
            cells = [clean_text("" if cell is None else str(cell)) for cell in row]
            if any(cells):
                rows.append(" | ".join(cells))
        if len(rows) > 1:
            parts.append("\n".join(rows))
    return "\n\n".join(parts)


def read_pptx(path: Path) -> str:
    ns = {
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    }
    parts = []
    with zipfile.ZipFile(path, "r") as archive:
        slide_names = sorted(
            [name for name in archive.namelist() if name.startswith("ppt/slides/slide") and name.endswith(".xml")]
        )
        for index, name in enumerate(slide_names, start=1):
            root = ET.fromstring(archive.read(name))
            texts = []
            for node in root.findall(".//a:t", ns):
                value = clean_text(node.text or "")
                if value:
                    texts.append(value)
            if texts:
                parts.append(f"## 第 {index} 页\n\n" + "\n".join(texts))
    return "\n\n".join(parts)


def read_text_like(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except Exception:
            continue
    return path.read_text(errors="ignore")


def main() -> int:
    if len(sys.argv) != 2:
        print(json.dumps({"error": "usage: extract_document.py <path>"}))
        return 1

    file_path = Path(sys.argv[1])
    suffix = file_path.suffix.lower()

    try:
        if suffix == ".pdf":
            text = read_pdf(file_path)
        elif suffix == ".docx":
            text = read_docx(file_path)
        elif suffix == ".xlsx":
            text = read_xlsx(file_path)
        elif suffix == ".pptx":
            text = read_pptx(file_path)
        elif suffix in {".md", ".markdown", ".txt", ".csv"}:
            text = read_text_like(file_path)
        else:
            raise ValueError(f"unsupported file type: {suffix}")

        text = clean_text(text)
        print(json.dumps({
            "ok": True,
            "path": str(file_path),
            "text": text,
            "chars": len(text),
        }, ensure_ascii=False))
        return 0
    except Exception as exc:
        print(json.dumps({
            "ok": False,
            "path": str(file_path),
            "error": str(exc),
        }, ensure_ascii=False))
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
